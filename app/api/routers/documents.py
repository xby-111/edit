"""
文档管理路由 - 统一使用 Service 层
"""
import io
import logging
import re
from typing import List, Optional
from urllib.parse import quote

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile, status
from fastapi.responses import Response

from app.core.security import get_current_user
from app. db.session import get_db
from app.schemas.comment import Comment, CommentCreate
from app.schemas. task import Task, TaskCreate, TaskUpdate
from app.schemas import (
    Document,
    DocumentCreate,
    DocumentUpdate,
    Template,
    TemplateCreate,
    TemplateUpdate,
)
from app.services. document_service import (
    create_document,
    create_document_version,
    create_template,
    delete_document,
    delete_template,
    get_document,
    get_document_versions,
    get_documents,
    get_folders,
    get_tags,
    get_template,
    get_templates,
    lock_document,
    search_documents,
    unlock_document,
    update_document,
    update_template,
    add_collaborator,
    batch_add_collaborators,
    remove_collaborator,
    get_collaborators,
    get_shared_documents,
    check_document_permission,
    is_document_owner,
    get_document_with_collaborators,
)
from app.services.audit_service import log_action
from app.services.settings_service import is_feature_enabled
from app.services. comment_service import create_comment, list_comments
from app.services.task_service import create_task, list_tasks, update_task

# PDF/Word 导入导出依赖
from xhtml2pdf import pisa
from docx import Document as DocxDocument
import pdfplumber

logger = logging.getLogger(__name__)

router = APIRouter(tags=["文档管理"])


# ==================== 工具函数 ====================

def htmlToMarkdown(html: str) -> str:
    """简单的 HTML 到 Markdown 转换"""
    if not html:
        return ""
    
    try:
        markdown = html
        # 标题
        markdown = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', markdown, flags=re. DOTALL)
        markdown = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', markdown, flags=re.DOTALL)
        markdown = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', markdown, flags=re.DOTALL)
        # 强调
        markdown = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', markdown, flags=re.DOTALL)
        markdown = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', markdown, flags=re. DOTALL)
        markdown = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', markdown, flags=re.DOTALL)
        markdown = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', markdown, flags=re.DOTALL)
        # 链接
        markdown = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', markdown, flags=re. DOTALL)
        # 图片
        markdown = re.sub(r'<img[^>]*src="([^"]*)"[^>]*alt="([^"]*)"[^>]*>', r'! [\2](\1)', markdown, flags=re.DOTALL)
        # 引用
        markdown = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'> \1', markdown, flags=re. DOTALL)
        # 代码
        markdown = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', markdown, flags=re.DOTALL)
        markdown = re.sub(r'<pre[^>]*>(.*?)</pre>', r'```\n\1\n```', markdown, flags=re. DOTALL)
        # 列表
        markdown = re.sub(r'<ul[^>]*>(.*? )</ul>', lambda m: re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', m.group(1), flags=re. DOTALL), markdown, flags=re. DOTALL)
        markdown = re. sub(r'<ol[^>]*>(.*?)</ol>', lambda m: re.sub(r'<li[^>]*>(.*?)</li>', r'1. \1', m.group(1), flags=re. DOTALL), markdown, flags=re. DOTALL)
        # 换行和段落
        markdown = re.sub(r'<br[^>]*>', '\n', markdown)
        markdown = re. sub(r'<p[^>]*>(.*?)</p>', r'\1\n', markdown, flags=re.DOTALL)
        # 移除所有剩余标签
        markdown = re.sub(r'<[^>]+>', '', markdown)
        return markdown. strip()
    except Exception:
        # 如果转换失败,返回空字符串
        return ""


def markdownToHtml(markdown: str) -> str:
    """简单的 Markdown 到 HTML 转换"""
    if not markdown:
        return ""
    
    try:
        html = markdown
        # 标题
        html = re.sub(r'^### (. +)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re. MULTILINE)
        html = re. sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        # 强调和斜体
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        # 链接
        html = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', html)
        # 图片
        html = re. sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1">', html)
        # 引用
        html = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', html, flags=re. MULTILINE)
        # 代码
        html = re.sub(r'```(.*?)```', r'<pre><code>\1</code></pre>', html, flags=re.DOTALL)
        html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)
        # 列表
        html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'(<li>.*</li>)', r'<ul>\1</ul>', html, flags=re.DOTALL)
        # 段落和换行
        html = re.sub(r'\n\n', '</p><p>', html)
        html = f'<p>{html}</p>'
        html = re.sub(r'<p></p>', '', html)
        html = re.sub(r'<p>(. +?)</p>\s*<p>', r'<p>\1</p>', html)
        return html
    except Exception:
        # 如果转换失败,返回原始内容包装在段落中
        return f"<p>{markdown}</p>"


# ==================== 文档列表 / 搜索 / 标签 / 文件夹相关路由 ====================

@router.get("/documents", response_model=List[Document], summary="获取文档列表", description="获取当前用户拥有的文档列表")
async def get_documents_endpoint(
    current_user = Depends(get_current_user), 
    db = Depends(get_db),
    skip: int = 0,
    limit: int = 100,
    folder: Optional[str] = None
):
    """获取当前用户拥有的文档列表"""
    documents = get_documents(db, current_user. id, skip=skip, limit=limit, folder=folder)
    return documents


@router.get("/documents/search", response_model=List[Document], summary="搜索文档", description="根据关键词、标签、日期等条件搜索文档")
async def search_documents_endpoint(
    keyword: Optional[str] = None,
    tags: Optional[str] = None,
    folder: Optional[str] = None,
    sort_by: str = "updated_at",
    order: str = "desc",
    created_from: Optional[str] = None,
    created_to: Optional[str] = None,
    updated_from: Optional[str] = None,
    updated_to: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """搜索文档"""
    documents = search_documents(
        db, current_user.id, keyword, tags, folder, sort_by, order,
        created_from, created_to, updated_from, updated_to, skip, limit
    )
    return documents


@router.get("/folders", response_model=List[str], summary="获取文件夹列表", description="获取用户的所有文件夹")
async def get_folders_endpoint(
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """获取文件夹列表"""
    folders = get_folders(db, current_user. id)
    return folders


@router. get("/tags", response_model=List[str], summary="获取标签列表", description="获取用户的所有标签")
async def get_tags_endpoint(
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """获取标签列表"""
    tags = get_tags(db, current_user. id)
    return tags


# ==================== 文档锁定/解锁相关路由 ====================

@router. post("/documents/{document_id}/lock", summary="锁定文档", description="锁定文档，禁止其他用户编辑")
async def lock_document_endpoint(
    document_id: int,
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """锁定文档"""
    # 检查文档权限（需要编辑权限才能锁定）
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_edit"]:
        raise HTTPException(status_code=403, detail="无编辑权限，无法锁定文档")
    
    # 获取文档信息检查锁定状态
    document = get_document_with_collaborators(db, document_id, current_user.id)
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    if document. get('is_locked'):
        raise HTTPException(status_code=400, detail="文档已被锁定")
    
    try:
        lock_document(db, document_id, current_user.id)
        return {"message": "文档已锁定"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error("锁定文档失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="锁定文档失败")


@router.post("/documents/{document_id}/unlock", summary="解锁文档", description="解锁文档，允许其他用户编辑")
async def unlock_document_endpoint(
    document_id: int,
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """解锁文档"""
    # 检查文档权限（需要编辑权限才能解锁）
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_edit"]:
        raise HTTPException(status_code=403, detail="无编辑权限，无法解锁文档")
    
    # 获取文档信息检查锁定状态
    document = get_document_with_collaborators(db, document_id, current_user.id)
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    if not document.get('is_locked'):
        raise HTTPException(status_code=400, detail="文档未被锁定")
    
    if document.get('locked_by') != current_user.id:
        raise HTTPException(status_code=403, detail="只有锁定者才能解锁文档")
    
    try:
        unlock_document(db, document_id, current_user.id)
        return {"message": "文档已解锁"}
    except HTTPException:
        raise
    except Exception as e:
        logger. error("解锁文档失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="解锁文档失败")


# ==================== 导出/导入相关路由 ====================

@router. get("/documents/{document_id}/export", summary="导出文档", description="将文档导出为指定格式")
async def export_document_endpoint(
    document_id: int,
    format: str = "html",
    current_user = Depends(get_current_user),
    db = Depends(get_db),
    request: Request = None,
):
    """导出文档"""
    if not is_feature_enabled(db, "feature. export. enabled", True):
        raise HTTPException(status_code=403, detail="导出功能已禁用")

    # 先检查文档是否存在且用户有权限
    document = get_document(db, document_id, current_user.id)
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    try:
        content = document.get('content', '')
        title = document.get('title', '文档')
        
        if format. lower() == 'markdown':
            # 简单的 HTML 到 Markdown 转换
            markdown_content = htmlToMarkdown(content)
            log_action(
                db,
                user_id=current_user.id,
                action="export. request",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"format": "markdown"},
            )
            ascii_name = f"document_{document_id}.md"
            cd = f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(ascii_name)}'
            return Response(
                content=markdown_content,
                media_type="text/markdown",
                headers={"Content-Disposition": cd}
            )
        elif format.lower() == 'html':
            log_action(
                db,
                user_id=current_user.id,
                action="export. request",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"format": "html"},
            )
            ascii_name = f"document_{document_id}.html"
            cd = f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(ascii_name)}'
            return Response(
                content=content,
                media_type="text/html",
                headers={"Content-Disposition": cd}
            )
        # 【注意】PDF 导出依赖于部署环境安装支持中文的字体（如 SimSun）。
        # 如果缺少相应字体，中文字符可能显示为乱码或方框。
        # 建议在服务器上安装 fonts-wqy-zenhei 或 fonts-noto-cjk 等字体包。
        elif format. lower() == 'pdf':
            # PDF 导出：使用 xhtml2pdf 将 HTML 转换为 PDF
            
            # 包装 HTML 内容，确保正确的编码和样式
            html_template = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{title}</title>
                <style>
                    body {{ font-family: SimSun, Arial, sans-serif; }}
                </style>
            </head>
            <body>
                {content}
            </body>
            </html>
            """
            
            pdf_buffer = io.BytesIO()
            pisa_status = pisa.CreatePDF(html_template, dest=pdf_buffer, encoding='utf-8')
            
            if pisa_status. err:
                raise HTTPException(status_code=500, detail="PDF 转换失败")
            
            pdf_buffer.seek(0)
            pdf_content = pdf_buffer.read()
            
            log_action(
                db,
                user_id=current_user.id,
                action="export. request",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"format": "pdf"},
            )
            ascii_name = f"document_{document_id}.pdf"
            cd = f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(ascii_name)}'
            return Response(
                content=pdf_content,
                media_type="application/pdf",
                headers={"Content-Disposition": cd}
            )
        # 【注意】Word 导出是基于 HTML 到 Markdown 的有损转换。
        # 复杂格式（如表格、嵌套列表、内联样式、图片等）会被丢失或简化。
        # 仅适用于纯文本和基本标题结构的文档。
        elif format.lower() == 'docx' or format.lower() == 'word':
            # Word 导出：使用 python-docx 将 HTML 转换为 DOCX
            
            doc = DocxDocument()
            
            # 将 HTML 转换为 Markdown 格式（用于提取结构化内容）
            text_content = htmlToMarkdown(content)
            
            # 按段落添加内容
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para. strip():
                    # 处理标题
                    if para. startswith('# '):
                        doc.add_heading(para[2:]. strip(), level=1)
                    elif para.startswith('## '):
                        doc.add_heading(para[3:].strip(), level=2)
                    elif para.startswith('### '):
                        doc. add_heading(para[4:].strip(), level=3)
                    else:
                        doc.add_paragraph(para.strip())
            
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            docx_content = docx_buffer.read()
            
            log_action(
                db,
                user_id=current_user.id,
                action="export.request",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"format": "docx"},
            )
            ascii_name = f"document_{document_id}. docx"
            cd = f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(ascii_name)}'
            return Response(
                content=docx_content,
                media_type="application/vnd. openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": cd}
            )
        else:
            raise HTTPException(status_code=400, detail="不支持的导出格式")
    except HTTPException:
        raise
    except Exception as e:
        logger.error("导出文档失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="导出文档失败")


@router.post("/documents/import", summary="导入文档", description="从上传的文件创建新文档")
async def import_document_endpoint(
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user = Depends(get_current_user), 
    db = Depends(get_db)
):
    """导入文档"""
    try:
        # 读取文件内容
        content = await file.read()
        
        # 根据文件类型处理内容
        if file.filename.endswith('.md') or file.filename. endswith('.txt'):
            # Markdown 或纯文本文件，直接使用
            text_content = content.decode('utf-8')
            html_content = markdownToHtml(text_content)
        elif file.filename.endswith('.html'):
            # HTML 文件，直接使用
            html_content = content.decode('utf-8')
        elif file.filename.endswith('.docx'):
            # Word 文档导入：使用 python-docx 读取
            
            docx_buffer = io.BytesIO(content)
            doc = DocxDocument(docx_buffer)
            
            # 提取段落内容并转换为 HTML
            html_parts = []
            for para in doc.paragraphs:
                text = para.text. strip()
                if text:
                    # 检查段落样式判断标题级别
                    if para.style.name. startswith('Heading 1'):
                        html_parts.append(f'<h1>{text}</h1>')
                    elif para.style. name.startswith('Heading 2'):
                        html_parts.append(f'<h2>{text}</h2>')
                    elif para.style.name. startswith('Heading 3'):
                        html_parts. append(f'<h3>{text}</h3>')
                    else:
                        html_parts. append(f'<p>{text}</p>')
            
            html_content = '\n'.join(html_parts)
        elif file.filename.endswith('.pdf'):
            # PDF 文档导入：使用 pdfplumber 提取文本
            
            pdf_buffer = io. BytesIO(content)
            html_parts = []
            
            with pdfplumber. open(pdf_buffer) as pdf:
                for page in pdf.pages:
                    text = page. extract_text()
                    if text:
                        # 按行分割并转换为段落
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                html_parts.append(f'<p>{line}</p>')
            
            html_content = '\n'. join(html_parts)
        else:
            # 其他文件类型，作为纯文本处理
            text_content = content.decode('utf-8')
            html_content = f"<p>{text_content. replace(chr(10), '<br>')}</p>"
        
        # 创建文档
        document_data = {
            'title': title,
            'content': html_content,
            'status': 'active'
        }
        
        new_document = create_document(db, document_data, current_user.id)
        
        return {
            "id": new_document['id'],
            "title": new_document['title'],
            "message": "文档导入成功"
        }
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="文件编码错误，请使用 UTF-8 编码的文件")
    except HTTPException:
        raise
    except Exception as e:
        logger.error("导入文档失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="导入文档失败")


# ==================== 评论与任务基础接口 ====================

@router.get("/documents/{document_id}/comments", response_model=List[Comment], summary="获取文档评论")
async def get_document_comments(
    document_id: int,
    current_user=Depends(get_current_user),
    db=Depends(get_db)
):
    """获取文档的评论列表"""
    # 权限检查：只需 can_view 权限
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_view"]:
        raise HTTPException(status_code=403, detail="文档不存在或无权访问")
    
    try:
        return list_comments(db, document_id)
    except Exception as e:
        logger.error("获取评论失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="获取评论失败")


@router.post("/documents/{document_id}/comments", response_model=Comment, summary="创建评论")
async def create_document_comment(
    document_id: int,
    comment_in: CommentCreate,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
    request: Request = None,
):
    """创建文档评论"""
    # 权限检查：只需 can_view 权限（通常允许协作者和查看者评论）
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_view"]:
        raise HTTPException(status_code=403, detail="文档不存在或无权访问")
        
    try:
        comment = create_comment(
            db,
            document_id,
            current_user.id,
            comment_in.content,
            comment_in.line_no,
            None,
            None,
            None,
        )
        # 记录操作日志
        try:
            log_action(
                db,
                user_id=current_user.id,
                action="comment.create",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"comment_id": getattr(comment, 'id', None) if hasattr(comment, 'id') else comment.get("id") if isinstance(comment, dict) else None},
            )
        except Exception:
            pass
        return comment
    except Exception as e:
        logger.error("创建评论失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="创建评论失败")


@router.get("/documents/{document_id}/tasks", response_model=List[Task], summary="获取文档任务")
async def get_document_tasks(
    document_id: int,
    current_user=Depends(get_current_user),
    db=Depends(get_db)
):
    """获取文档的任务列表"""
    # 权限检查：只需 can_view 权限
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_view"]:
        raise HTTPException(status_code=403, detail="文档不存在或无权访问")
        
    try:
        return list_tasks(db, document_id)
    except Exception as e:
        logger.error("获取任务失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="获取任务失败")


@router.post("/documents/{document_id}/tasks", response_model=Task, summary="创建任务")
async def create_document_task(
    document_id: int,
    task_in: TaskCreate,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
    request: Request = None,
):
    """创建文档任务"""
    # 权限检查：必须有 can_edit 权限才能创建任务
    permission = check_document_permission(db, document_id, current_user.id)
    if not permission["can_edit"]:
        raise HTTPException(status_code=403, detail="无编辑权限，无法创建任务")

    try:
        # 处理任务输入数据
        assignee_id = getattr(task_in, 'assignee_id', None) or getattr(task_in, 'assigned_to', None)
        due = getattr(task_in, "due_at", None) or getattr(task_in, "due_date", None) or getattr(task_in, "deadline", None)
        
        task = create_task(
            db,
            document_id,
            current_user.id,
            task_in.title,
            task_in.description,
            assignee_id,
            due,
        )
        
        # 记录操作日志
        try:
            log_action(
                db,
                user_id=current_user.id,
                action="task.create",
                resource_type="document",
                resource_id=document_id,
                request=request,
                meta={"task_id": task.get("id") if isinstance(task, dict) else getattr(task, "id", None)},
            )
        except Exception:
            pass
            
        return task
    except Exception as e:
        logger.error("创建任务失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="创建任务失败")


# ==================== 文档 CRUD 相关路由 ====================

@router.post("/documents", response_model=Document, status_code=status.HTTP_201_CREATED, summary="创建文档", description="创建新的协作文档")
async def create_document_endpoint(
    document: DocumentCreate,
    current_user = Depends(get_current_user),
    db = Depends(get_db),
    request: Request = None,
):
    """创建文档"""
    try:
        new_document = create_document(db, document, current_user.id)
        try:
            log_action(
                db,
                user_id=current_user.id,
                action="document.create",
                resource_type="document",
                resource_id=new_document. get("id"),
                request=request,
            )
        except Exception:
            pass

        return Document(
            id=new_document['id'],
            owner_id=new_document['owner_id'],
            title=new_document['title'],
            content=new_document['content'],
            status=new_document['status'],
            folder_name=new_document. get('folder_name'),
            tags=new_document. get('tags'),
            is_locked=new_document.get('is_locked', False),
            created_at=new_document['created_at'],
            updated_at=new_document['updated_at']
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("创建文档失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="创建文档失败")


@router.get("/documents/shared", response_model=List[Document], summary="获取共享文档列表", description="获取当前用户作为协作者的文档列表")
async def get_shared_documents_endpoint(
    current_user = Depends(get_current_user),
    db = Depends(get_db),
    skip: int = 0,
    limit: int = 100
):
    """获取当前用户作为协作者的文档列表"""
    documents = get_shared_documents(db, current_user.id, skip=skip, limit=limit)
    return documents